# read and modify excel spreadsheets

import openpyxl

workbook = openpyxl.load_workbook("example.xlsx")

print(type(workbook))
print(workbook.get_sheet_names())

sheet = workbook.get_sheet_by_name("Sheet1")
# print(sheet)

print(sheet['A1'])
cell = sheet['A1']
print(str(cell.value))

print(sheet.cell(row = 1, column = 2))

# Editing Excel Spreadsheets

sheet['A1'] = "Apple Station"

# save progress
workbook.save("example.xlsx")

# create new worksheet

sheet122 = workbook.create_sheet()
sheet122 = "MY New Sheet"

workbook.save("example.xlsx")

# PDF FILE => Binary File

import PyPDF2 as pdfReader

pdfFile = open("14.pdf","rb")# read binary mode

reader = pdfReader.PdfFileReader(pdfFile)

# number of pages
print(reader.numPages)

pages = reader.getPage(0)
pages = pages.extractText()
print(pages)

# it doesnt work in mine

# pdfFile2 = open("14.pdf","rb")

# reader2 = pdfReader.PdfFileReader(pdfFile2)
# writer = pdfReader.PdfFileWriter()

# for pageNum in range(reader.numPages):
#     page = reader.getPage(pageNum)
#     writer.addPage(page)

# for pageNum in range(reader2.numPages):
#     page = reader2.getPage(pageNum)
#     writer.addPage(page)

# outputFile = open("14-double.pdf","wb")
# writer.write(outputFile)

# outputFile.close()
# pdfFile.close()
# pdfFile2.close()

# CREATE AND EDIT Word Documents

import docx

d = docx.Document("14.docx")

# the texts
print (d.paragraphs)

secondParagraph = d.paragraphs[3]
print(secondParagraph.text) 

secondParagraph.bold = True
secondParagraph.italic = True
secondParagraph.underline = True
# we can also check if bold, italic, etc using .bold or .italic -> returns true

# d.save("14.docx")

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)

print(getText("14.docx"))

''' 
 Open with docx.Document()
 Paragraphs are composed of runs
we can set each runs to italic, bold, or underline
word files can be created by calling add_paragraphs() and add run()
to append text
```
